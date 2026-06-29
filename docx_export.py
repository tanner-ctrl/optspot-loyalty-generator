"""docx_export.py — Google Docs-compatible Word export for loyalty drafts.

Produces a .docx the customer can upload to Google Drive and open in Google
Docs to leave comments and edits. Uses Word heading styles so Google Docs
auto-builds a table of contents on conversion.
"""

import io

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


NAVY = RGBColor(0x26, 0x40, 0x78)
GRAY = RGBColor(0x88, 0x88, 0x88)
BODY_FONT = "Calibri"


# ── Incentive helpers (mirror pdf_export Incentives Summary) ────────────────────

def _incentive_items(context: dict) -> list:
    items = []
    prog_type = context.get("program_type", "visit-based")

    if prog_type == "visit-based":
        tiers = context.get("tiers", [])
        if tiers:
            first = tiers[0]
            visits = first.get("visits", "")
            reward = first.get("reward", "")
            if visits and reward:
                items.append(f"Earn a {reward} after {visits} visits")
    else:
        for pkg in context.get("wash_packages", []):
            name = pkg.get("name", "")
            earn = pkg.get("earn_points", "")
            redeem = pkg.get("redeem_cost", "")
            if name and earn and redeem:
                items.append(f"{name}: earn {earn} pt per wash, redeem at {redeem} pts")

    if context.get("hpo_enabled", True):
        offer = (context.get("hpo_membership_offer") or "").strip()
        if offer:
            items.append(f"HPO Membership Offer: {offer}")
        if context.get("hpo_timeframe_days"):
            items.append(f"HPO Timeframe: {context['hpo_timeframe_days']} days")
        if context.get("hpo_min_visits"):
            items.append(f"HPO Minimum Visits: {context['hpo_min_visits']}")
        if context.get("hpo_max_checkins"):
            items.append(f"HPO Maximum Check-ins: {context['hpo_max_checkins']}")
        hpo_exec = context.get("hpo_execution")
        if hpo_exec:
            labels = {
                "link": "Link to ecommerce ([Ecomm LINK])",
                "onsite": "On-site redemption (in-store only)",
                "redeem": "Hard offer (~redeem~)",
            }
            items.append(f"HPO Execution: {labels.get(hpo_exec, hpo_exec)}")

    for ae in context.get("auto_engage", []):
        if ae.get("type", "offer") == "offer" and (ae.get("offer") or "").strip():
            items.append(f"{ae['days']} Day Offer: {ae['offer']}")

    return items


def _code_items(context: dict) -> list:
    items = []
    if context.get("uses_redemption_codes"):
        fields = [
            ("codes_location", "Codes managed at"),
            ("code_format_example", "Code format"),
            ("codes_provider", "Codes provided by"),
            ("code_delivery_method", "Delivery method"),
        ]
        for key, label in fields:
            val = (context.get(key) or "").strip()
            if val:
                items.append(f"{label}: {val}")
    return items


# ── Low-level paragraph helpers ─────────────────────────────────────────────────

def _set_base_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(11)


def _fire_desc(p):
    """Style an italic gray 'when it fires' description line."""
    p.runs[0].italic = True
    p.runs[0].font.color.rgb = GRAY
    p.runs[0].font.size = Pt(9.5)


def _add_message_block(doc, option_label: str, text: str) -> None:
    """Add an 'Option N (MODE)' label, the message text, then 2 blank lines."""
    lbl = doc.add_paragraph()
    run = lbl.add_run(option_label)
    run.bold = True
    run.font.color.rgb = NAVY
    run.font.size = Pt(10.5)

    body = doc.add_paragraph(text)
    body.paragraph_format.space_after = Pt(2)

    # Two blank lines so reviewers have room to comment in Google Docs.
    doc.add_paragraph("")
    doc.add_paragraph("")


# ── Section rendering ───────────────────────────────────────────────────────────

# (section_name, message description, getter) — getter pulls the message(s) from
# messages_data. Single-message sections return a dict; multi return a list.
_SECTION_DEFS = [
    ("Welcome", "Sent immediately when a customer joins the loyalty program.", "welcome"),
    ("Visit Tracked", "Sent right after each wash to confirm the visit and show progress.", "tracked"),
    ("Progress", "Sent after a wash with a link for the customer to check loyalty progress.", "progress"),
    ("Rewards", "Sent when a customer hits a reward milestone.", "rewards"),
    ("Auto-Engage", "Sent when a customer hasn't visited in a while, to bring them back.", "auto_engage"),
    ("Hot Prospect", "Sent to high-frequency customers as a VIP thank-you and upgrade incentive.", "hot_prospect"),
]


def _render_section(doc, section_name: str, desc: str, payload) -> None:
    """Render one message-type section. payload is a dict, a list, or None.

    Returns without adding anything when the section is disabled/empty.
    """
    if not payload:
        return

    # Normalize to a list of message dicts.
    entries = payload if isinstance(payload, list) else [payload]
    entries = [m for m in entries if m and m.get("text")]
    if not entries:
        return

    doc.add_heading(section_name, level=1)
    desc_p = doc.add_paragraph(desc)
    _fire_desc(desc_p)

    multi = len(entries) > 1
    for m in entries:
        if multi:
            doc.add_heading(m.get("label", section_name), level=2)
        mode = m.get("mode", "SMS")
        _add_message_block(doc, f"Option 1 ({mode})", m.get("text", ""))


# ── Public API ───────────────────────────────────────────────────────────────────

def build_loyalty_docx(context: dict, messages_data: dict) -> bytes:
    """Build the loyalty draft as a .docx and return it as bytes."""
    doc = Document()
    _set_base_font(doc)

    # 1-inch margins all sides.
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    car_wash = context.get("car_wash_name") or "Car Wash"

    # ── Title + subtitle ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    t_run = title.add_run(f"{car_wash} — Loyalty Message Draft")
    t_run.bold = True
    t_run.font.size = Pt(22)
    t_run.font.color.rgb = NAVY

    subtitle = doc.add_paragraph()
    s_run = subtitle.add_run("Prepared by OptSpot. Review, edit, and return for finalization.")
    s_run.italic = True
    s_run.font.size = Pt(11)
    s_run.font.color.rgb = GRAY
    doc.add_paragraph("")

    # ── Program Overview ──
    doc.add_heading("Program Overview", level=1)
    prog_type = context.get("program_type", "visit-based")
    prog_label = "Visit-based" if prog_type == "visit-based" else "Points-based"
    signup_on = bool(context.get("signup_reward_enabled") and context.get("signup_reward"))
    signup_txt = (
        f"On — {context.get('signup_reward')}" if signup_on else "Off"
    )
    view_state = context.get("view_mode", "Tabs")
    for line in [
        f"Program type: {prog_label}",
        f"Signup reward: {signup_txt}",
        f"View mode: {view_state}",
    ]:
        doc.add_paragraph(line, style="List Bullet")
    doc.add_paragraph("")

    # ── Program Incentives ──
    doc.add_heading("Program Incentives", level=1)
    items = _incentive_items(context)
    code_items = _code_items(context)
    if items:
        for it in items:
            doc.add_paragraph(it, style="List Bullet")
    if code_items:
        sub = doc.add_heading("Redemption Codes", level=2)
        for it in code_items:
            doc.add_paragraph(it, style="List Bullet")
    if not items and not code_items:
        doc.add_paragraph("No incentives configured.")
    doc.add_paragraph("")

    # ── Message sections ──
    for section_name, desc, key in _SECTION_DEFS:
        _render_section(doc, section_name, desc, messages_data.get(key))

    # ── Footer ──
    footer_para = doc.sections[0].footer.paragraphs[0]
    footer_para.text = "OptSpot Internal Draft — for client review only."
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_para.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = GRAY

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
